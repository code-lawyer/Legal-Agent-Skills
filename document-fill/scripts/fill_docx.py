#!/usr/bin/env python3
"""fill_docx.py — 把填充计划渲染进模板。
docx 优先：python-docx 就地替换 {{slot_id}} 占位符。python-docx 缺失或模板非 docx 时，
降级为 markdown 填好稿并返回 degraded=True，不静默失败。
本 skill 只渲染、不做法律判断；gap/待起草/存疑用可见标记，绝不伪造。"""
import json, sys
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def render_value(item):
    st = item.get("status")
    label = item.get("slot_label", item.get("slot_id", ""))
    if st == "gap":
        return f"【缺口：{label}】"
    if st == "pending_drafting":
        return f"【待起草：{label}】"
    if st == "ambiguous":
        vals = " / ".join(str(c.get("value", "")) for c in (item.get("candidates") or []))
        return f"【存疑·并列待裁：{vals}】"
    value = item.get("value")
    return str(value) if value is not None else ""

def _index(plan):
    return {it.get("slot_id"): it for it in plan}

def fill_markdown(template_text, plan):
    out = template_text
    for sid, it in _index(plan).items():
        out = out.replace("{{" + sid + "}}", render_value(it))
    return out

def _degraded_note(plan):
    lines = ["# 降级填充稿（原模板无法就地渲染）", ""]
    for it in plan:
        lines.append(f"- {it.get('slot_label', it.get('slot_id', ''))}: {render_value(it)}")
    return "\n".join(lines)

def _replace_in_paragraph(para, token, value):
    if token not in para.text:
        return
    new_text = para.text.replace(token, value)
    for r in list(para.runs):
        r.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

def fill_docx(template_path, plan, out_path):
    tpl = Path(template_path)
    if not DOCX_AVAILABLE or tpl.suffix.lower() != ".docx":
        if tpl.suffix.lower() in {".md", ".txt"}:
            md = fill_markdown(tpl.read_text(encoding="utf-8"), plan)
        else:
            md = _degraded_note(plan)
        md_out = Path(out_path).with_suffix(".md")
        md_out.write_text(md, encoding="utf-8")
        return {"degraded": True, "output": str(md_out)}
    doc = Document(str(tpl))
    for para in doc.paragraphs:
        for sid, it in _index(plan).items():
            token = "{{" + sid + "}}"
            if token in para.text:
                _replace_in_paragraph(para, token, render_value(it))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for sid, it in _index(plan).items():
                        token = "{{" + sid + "}}"
                        if token in para.text:
                            _replace_in_paragraph(para, token, render_value(it))
    doc.save(out_path)
    return {"degraded": False, "output": str(out_path)}

def main(argv):
    if len(argv) != 4:
        print("用法: python fill_docx.py <模板> <fill_plan.json> <输出.docx>")
        return 2
    with open(argv[2], encoding="utf-8") as f:
        plan = json.load(f)
    res = fill_docx(argv[1], plan, argv[3])
    print(("⚠️ 降级为 markdown：" if res["degraded"] else "✅ 已渲染：") + res["output"])
    return 0

if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    sys.exit(main(sys.argv))
