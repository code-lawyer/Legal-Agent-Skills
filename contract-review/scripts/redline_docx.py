"""redline_docx.py

Turn a redline plan into a Word document with real tracked changes
(w:ins / w:del) and Word comments (批注), starting from an original
contract (.docx or .txt/.md).

A redline plan is a JSON array of 修订项 (redline-item) dicts, one per
change. Fields consumed per item:
  clause, action (insert|delete|replace|comment), anchor_text,
  tracked_changes ([-删-]/[+增+] markup, optional), clean_version,
  reason, legal_basis, negotiation_point (optional), risk_level.
The agent generates this plan from its problem cards / 建议修订文本;
this script only renders it — it does not make legal judgements.

python-docx has no native *tracked-changes* API, so w:ins/w:del elements
are built directly via docx.oxml.OxmlElement + docx.oxml.ns.qn. It does
have a native *comments* API as of 1.2.0 (Document.add_comment); this
module uses that when available and falls back to hand-built oxml
(a minimal word/comments.xml part + commentRangeStart/End/Reference) when
it is not.

Degradation: if python-docx is not importable at all (DOCX_AVAILABLE is
False), apply_redline cannot produce a .docx. Instead it reads the plan
and writes a markdown side-by-side comparison table to the output path's
`.md` sibling, returning `degraded=True` rather than raising.

Anchor/span caveat: matching is paragraph-scoped -- `anchor_text` must be
locatable within a single paragraph's text. Paragraph text outside the
matched anchor span (prefix/suffix in the same paragraph) is preserved
as plain text rather than being dropped, but its original character-level
formatting (bold/italic/etc. rPr) is NOT preserved in the rebuilt
paragraph (v1 limitation -- fine for plain-text contracts; formatted
.docx deliverables need a manual review pass before sign-off).
"""

import argparse
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.run import Run as _DocxRun

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Matches [-deleted text-] or [+inserted text+] segments in a
# tracked_changes string; anything else is plain "keep" text.
_TRACKED_RE = re.compile(r"\[-(.*?)-\]|\[\+(.*?)\+\]")


def load_plan(path):
    """Load the redline plan JSON (a list of redline item dicts)."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def open_or_wrap(input_path):
    """Open a .docx directly, or wrap a .txt/.md file into a new Document
    (one paragraph per line)."""
    lower = str(input_path).lower()
    if lower.endswith(".docx"):
        return Document(input_path)
    doc = Document()
    with open(input_path, "r", encoding="utf-8") as f:
        for line in f.read().splitlines():
            doc.add_paragraph(line)
    return doc


def parse_tracked(s):
    """Parse a tracked_changes string into a list of
    ("keep"|"del"|"ins", text) segments.

    Example: "A[-B-][+C+]D" -> [("keep","A"),("del","B"),("ins","C"),("keep","D")]
    """
    segments = []
    pos = 0
    for m in _TRACKED_RE.finditer(s):
        if m.start() > pos:
            segments.append(("keep", s[pos:m.start()]))
        if m.group(1) is not None:
            segments.append(("del", m.group(1)))
        else:
            segments.append(("ins", m.group(2)))
        pos = m.end()
    if pos < len(s):
        segments.append(("keep", s[pos:]))
    return segments


def locate_paragraph(doc, anchor):
    """Return the first paragraph whose text contains `anchor`, or None."""
    if not anchor:
        return None
    for para in doc.paragraphs:
        if anchor in para.text:
            return para
    return None


def _next_id_factory():
    counter = {"n": 0}

    def _next():
        counter["n"] += 1
        return str(counter["n"])

    return _next


def _iso_now():
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _mk_text_run(text, text_tag="w:t"):
    """Build a plain w:r run wrapping a single text element (`w:t` for
    normal runs, `w:delText` for the run inside a w:del)."""
    r = OxmlElement("w:r")
    t = OxmlElement(text_tag)
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _mk_tracked(wrapper_tag, text_tag, text, author, date, wid):
    """Build a tracked-change wrapper (w:ins/w:del) with id/author/date
    around a single text run."""
    wrapper = OxmlElement(wrapper_tag)
    wrapper.set(qn("w:id"), wid)
    wrapper.set(qn("w:author"), author)
    wrapper.set(qn("w:date"), date)
    wrapper.append(_mk_text_run(text, text_tag))
    return wrapper


def _mk_ins(text, author, date, wid):
    """Build a w:ins element containing a w:r/w:t run with `text`."""
    return _mk_tracked("w:ins", "w:t", text, author, date, wid)


def _mk_del(text, author, date, wid):
    """Build a w:del element containing a w:r/w:delText run with `text`."""
    return _mk_tracked("w:del", "w:delText", text, author, date, wid)


def _mk_keep_run(text):
    """Build a plain (untracked) w:r/w:t run with `text`."""
    return _mk_text_run(text)


def _merge_keep_segments(segments):
    """Collapse consecutive ("keep", text) entries into one, so a
    paragraph prefix/suffix stitched around a tracked-change span ends up
    as a single contiguous run instead of several adjacent (but visually
    identical) keep runs. Drops empty-text entries."""
    merged = []
    for kind, text in segments:
        if not text:
            continue
        if kind == "keep" and merged and merged[-1][0] == "keep":
            merged[-1] = ("keep", merged[-1][1] + text)
        else:
            merged.append((kind, text))
    return merged


def _coalesce_keep_parts(parts):
    """Collapse consecutive ("keep", text) parts into one, so plain-text
    fragments stitched around tracked-change nodes end up as a single
    contiguous run. ("el", element) parts (preserved or newly-built
    w:ins/w:del) are pass-through boundaries and never merge across.
    Drops empty keep text."""
    merged = []
    for kind, val in parts:
        if kind == "keep":
            if not val:
                continue
            if merged and merged[-1][0] == "keep":
                merged[-1] = ("keep", merged[-1][1] + val)
                continue
        merged.append((kind, val))
    return merged


def _run_text(r):
    """Concatenate the visible text (w:t children) of a plain w:r run."""
    return "".join(t.text or "" for t in r.findall(qn("w:t")))


def _split_run_text(text, start, a_start, a_end):
    """For a plain run occupying plain-text offsets [start, start+len(text)),
    return (before, after): the sub-text falling before a_start and after
    a_end respectively. Text inside [a_start, a_end) is the anchor span
    itself and is dropped (the middle segments replace it)."""
    end = start + len(text)
    before = text[: min(a_start, end) - start] if start < a_start else ""
    after = text[max(a_end, start) - start:] if end > a_end else ""
    return before, after


def _rewrite_paragraph_span(para, anchor_text, middle_segments, author, date, next_id):
    """Replace the `anchor_text` span in `para` with `middle_segments`
    (a list of (kind, text)) rendered as keep/del/ins runs, while PRESERVING
    any pre-existing tracked-change nodes (w:ins/w:del from earlier redline
    items applied to the same paragraph) in their original document order.

    The anchor is located within the paragraph's plain-run text only — the
    same text locate_paragraph / para.text sees — and existing w:ins/w:del
    are treated as opaque and kept in place. This lets multiple redline
    items touch one paragraph without an earlier change migrating to the
    paragraph start (the failure mode of the previous whole-paragraph
    rebuild, which orphaned prior w:ins/w:del at the front).

    Returns (changes_count, anchor_info). anchor_info is
    (first_run_el, last_run_el) spanning the del/ins runs actually produced
    (for comment-range anchoring), or None if none were produced.

    Note (v1 limitation, unchanged): prefix/suffix plain text is preserved
    but original character-level rPr formatting is not carried over."""
    p = para._p
    # Snapshot run-level children in order: plain runs (splittable) and
    # existing tracked-change wrappers (preserved opaque).
    atoms = []
    for child in list(p):
        if child.tag == qn("w:r"):
            atoms.append(("keep", child))
        elif child.tag in (qn("w:ins"), qn("w:del")):
            atoms.append(("preserve", child))

    keep_text = "".join(_run_text(el) for kind, el in atoms if kind == "keep")
    idx = keep_text.find(anchor_text) if anchor_text else -1
    if idx < 0:
        # Anchor no longer locatable in plain text (e.g. it now overlaps an
        # existing tracked change); leave the paragraph untouched.
        return 0, None
    a_start, a_end = idx, idx + len(anchor_text)

    # Detach all run-level children (keep w:pPr and other paragraph metadata).
    for _, el in atoms:
        p.remove(el)

    # Build an ordered list of output "parts": ("keep", text) for plain text
    # and ("el", element) for preserved / newly-created tracked-change nodes.
    before_parts = []
    after_parts = []
    kpos = 0
    for kind, el in atoms:
        if kind == "preserve":
            # Bucket an existing tracked change by its position in plain text.
            (before_parts if kpos < a_end else after_parts).append(("el", el))
            continue
        text = _run_text(el)
        before, after = _split_run_text(text, kpos, a_start, a_end)
        kpos += len(text)
        if before:
            before_parts.append(("keep", before))
        if after:
            after_parts.append(("keep", after))

    # Render the middle segments into the vacated anchor position.
    changes = 0
    first_run_el = None
    last_run_el = None
    middle_parts = []
    for kind, text in _merge_keep_segments(middle_segments):
        if kind == "keep":
            middle_parts.append(("keep", text))
            continue
        el = _mk_del(text, author, date, next_id()) if kind == "del" \
            else _mk_ins(text, author, date, next_id())
        middle_parts.append(("el", el))
        changes += 1
        run_el = el.find(qn("w:r"))
        if first_run_el is None:
            first_run_el = run_el
        last_run_el = run_el

    # Coalesce consecutive plain-text parts into single runs (never merging
    # across a preserved tracked-change node), then materialize in order.
    for kind, val in _coalesce_keep_parts([*before_parts, *middle_parts, *after_parts]):
        p.append(_mk_keep_run(val) if kind == "keep" else val)

    anchor_info = (first_run_el, last_run_el) if first_run_el is not None else None
    return changes, anchor_info


def _apply_replace(doc, item, author, date, next_id):
    """Rewrite the anchor span as tracked del/ins. Uses tracked_changes'
    [-x-]/[+y+] markup when present, else a whole-anchor del + clean_version
    ins. Prefix/suffix text in the paragraph is preserved."""
    anchor_text = item.get("anchor_text", "")
    para = locate_paragraph(doc, anchor_text)
    if para is None:
        return 0, None
    tracked = item.get("tracked_changes") or ""
    if tracked:
        middle = parse_tracked(tracked)
    else:
        middle = [("del", anchor_text), ("ins", item.get("clean_version", ""))]
    return _rewrite_paragraph_span(para, anchor_text, middle, author, date, next_id)


def _apply_delete(doc, item, author, date, next_id):
    """Delete only the anchor_text span within the paragraph, preserving
    any prefix/suffix text in the same paragraph as plain keep runs."""
    anchor_text = item.get("anchor_text", "")
    para = locate_paragraph(doc, anchor_text)
    if para is None:
        return 0, None
    return _rewrite_paragraph_span(
        para, anchor_text, [("del", anchor_text)], author, date, next_id
    )


def _apply_insert(doc, item, author, date, next_id):
    anchor_para = locate_paragraph(doc, item.get("anchor_text", ""))
    clean = item.get("clean_version", "")
    new_p = OxmlElement("w:p")
    el = _mk_ins(clean, author, date, next_id())
    new_p.append(el)
    if anchor_para is not None:
        anchor_para._p.addnext(new_p)
    else:
        doc.element.body.append(new_p)
    run_el = el.find(qn("w:r"))
    return 1, (run_el, run_el)


def _comment_text(item):
    """Build the 批注 body text for a redline item: reason｜依据:legal_basis
    (+ ｜谈判:negotiation_point if present)."""
    reason = item.get("reason", "")
    legal_basis = item.get("legal_basis", "")
    text = f"{reason}｜依据:{legal_basis}"
    negotiation_point = item.get("negotiation_point")
    if negotiation_point:
        text += f"｜谈判:{negotiation_point}"
    return text


def _get_or_create_fallback_comments_part(doc):
    """Lazily create (or return the existing) word/comments.xml part built
    by hand, for python-docx installs that lack a native Document.add_comment
    (e.g. older than 1.2.0). Cached on the Document instance so repeated
    comments in the same document share one part."""
    part = getattr(doc, "_redline_fallback_comments_part", None)
    if part is not None:
        return part

    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    from docx.opc.constants import RELATIONSHIP_TYPE as RT, CONTENT_TYPE as CT
    from docx.oxml import parse_xml
    from lxml import etree

    class _CommentsFallbackPart(Part):
        def __init__(self, package):
            super().__init__(PackURI("/word/comments.xml"), CT.WML_COMMENTS, None, package)
            self.element = parse_xml(
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/'
                'wordprocessingml/2006/main"/>'
            )

        @property
        def blob(self):
            return etree.tostring(
                self.element, xml_declaration=True, encoding="UTF-8", standalone=True
            )

    part = _CommentsFallbackPart(doc.part.package)
    doc.part.relate_to(part, RT.COMMENTS)
    doc._redline_fallback_comments_part = part
    doc._redline_fallback_next_id = 0
    return part


def _add_comment_oxml_fallback(doc, first_el, last_el, text, author, initials):
    """Build commentRangeStart/commentRangeEnd/commentReference by hand and
    append a <w:comment> to a hand-built word/comments.xml part. Used only
    when the installed python-docx has no native add_comment API."""
    part = _get_or_create_fallback_comments_part(doc)
    comment_id = doc._redline_fallback_next_id
    doc._redline_fallback_next_id += 1

    comment_el = OxmlElement("w:comment")
    comment_el.set(qn("w:id"), str(comment_id))
    comment_el.set(qn("w:author"), author)
    comment_el.set(qn("w:initials"), initials or "")
    p_el = OxmlElement("w:p")
    p_el.append(_mk_text_run(text))
    comment_el.append(p_el)
    part.element.append(comment_el)

    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(comment_id))
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))
    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    ref_run.append(ref)

    first_el.addprevious(start)
    last_el.addnext(ref_run)
    last_el.addnext(end)
    return True


def _add_comment(doc, first_el, last_el, text, author, initials):
    """Attach a Word comment spanning first_el..last_el (w:r elements).
    Prefers python-docx 1.2.0's native Document.add_comment; falls back to
    hand-built oxml when that API is unavailable."""
    if hasattr(doc, "add_comment"):
        first_run = _DocxRun(first_el, None)
        last_run = _DocxRun(last_el, None)
        doc.add_comment([first_run, last_run], text=text, author=author, initials=initials)
        return True
    return _add_comment_oxml_fallback(doc, first_el, last_el, text, author, initials)


def _degrade_to_markdown(plan_path, output_path):
    """Write a markdown 对照稿 (side-by-side comparison table) for the plan
    to output_path's .md sibling. Used when python-docx is unavailable."""
    plan = load_plan(plan_path)

    def esc(value):
        return str(value or "").replace("|", "\\|").replace("\n", " ")

    lines = [
        "| 条款 | 动作 | 原文 anchor | 改写 clean_version | 理由 | 依据 |",
        "| --- | --- | --- | --- | --- | --- |",
    ]
    for item in plan:
        lines.append(
            "| {clause} | {action} | {anchor} | {clean} | {reason} | {basis} |".format(
                clause=esc(item.get("clause")),
                action=esc(item.get("action")),
                anchor=esc(item.get("anchor_text")),
                clean=esc(item.get("clean_version")),
                reason=esc(item.get("reason")),
                basis=esc(item.get("legal_basis")),
            )
        )

    md_path = Path(output_path).with_suffix(".md")
    md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return len(plan)


def apply_redline(input_path, plan_path, output_path, author="Reviewer"):
    """Apply a redline plan to an input contract, producing a .docx with
    Word tracked changes (w:ins/w:del) and Word comments (批注) carrying the
    reason/legal_basis/negotiation_point for each change.

    If python-docx is unavailable (DOCX_AVAILABLE is False), degrades to
    writing a markdown 对照稿 at output_path's .md sibling instead of
    raising.

    Returns a stats dict: {"changes": int, "comments": int,
    "degraded": bool, "skipped_comment": int, "warnings": [str, ...]}.
    """
    if not DOCX_AVAILABLE:
        n = _degrade_to_markdown(plan_path, output_path)
        return {
            "changes": 0,
            "comments": 0,
            "degraded": True,
            "skipped_comment": 0,
            "warnings": [f"python-docx 不可用,已降级为 markdown 对照稿({n} 条)"],
        }

    doc = open_or_wrap(input_path)
    plan = load_plan(plan_path)
    next_id = _next_id_factory()
    date = _iso_now()
    initials = (author[:2] if author else "RV") or "RV"

    changes = 0
    comments = 0
    skipped_comment = 0
    warnings = []

    for item in plan:
        action = item.get("action")
        clause = item.get("clause", "?")
        anchor_info = None

        if action == "replace":
            n, anchor_info = _apply_replace(doc, item, author, date, next_id)
            changes += n
        elif action == "delete":
            n, anchor_info = _apply_delete(doc, item, author, date, next_id)
            changes += n
        elif action == "insert":
            n, anchor_info = _apply_insert(doc, item, author, date, next_id)
            changes += n
        elif action == "comment":
            para = locate_paragraph(doc, item.get("anchor_text", ""))
            if para is not None and para.runs:
                anchor_info = (para.runs[0]._r, para.runs[-1]._r)
        else:
            # Unknown/unsupported action: count + warn rather than a
            # silent no-op, so callers/CLI users notice a plan typo.
            skipped_comment += 1
            warnings.append(f"未知动作 '{action}'(条款:{clause}),已跳过")
            continue

        if anchor_info is not None:
            first_el, last_el = anchor_info
            _add_comment(doc, first_el, last_el, _comment_text(item), author, initials)
            comments += 1
        else:
            skipped_comment += 1
            warnings.append(f"未找到锚点文本,批注跳过(条款:{clause})")

    doc.save(output_path)

    return {
        "changes": changes,
        "comments": comments,
        "degraded": False,
        "skipped_comment": skipped_comment,
        "warnings": warnings,
    }


def _build_arg_parser():
    parser = argparse.ArgumentParser(
        description="Apply a clause-redlining plan as Word tracked changes + comments "
        "(or a markdown fallback if python-docx is unavailable)."
    )
    parser.add_argument("--input", required=True, help="Original contract (.docx/.txt/.md)")
    parser.add_argument("--plan", required=True, help="Redline plan JSON path")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--author", default="Reviewer", help="Tracked-change/comment author name")
    return parser


def main():
    # Windows consoles often default stdout/stderr to a legacy codepage
    # (e.g. cp1252) that cannot encode the CJK text in reasons/warnings;
    # force UTF-8 so `print` doesn't crash on non-ASCII output.
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except (AttributeError, ValueError):
            pass

    parser = _build_arg_parser()
    args = parser.parse_args()
    result = apply_redline(args.input, args.plan, args.output, author=args.author)

    for warning in result.get("warnings", []):
        print(f"warning: {warning}", file=sys.stderr)

    if result.get("degraded"):
        md_path = Path(args.output).with_suffix(".md")
        print(
            f"note: python-docx unavailable, wrote markdown 对照稿 to {md_path} instead of .docx",
            file=sys.stderr,
        )

    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
