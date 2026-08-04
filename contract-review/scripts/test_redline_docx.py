import zipfile
from docx import Document
import redline_docx

def _make_input(tmp_path):
    p = tmp_path / "in.docx"
    d = Document(); d.add_paragraph("乙方离职后两年内不得从事与甲方相同或类似业务。")
    d.save(str(p)); return p

def test_replace_produces_ins_del(tmp_path):
    inp = _make_input(tmp_path)
    plan = tmp_path / "plan.json"
    plan.write_text('[{"clause":"第8条","action":"replace",'
        '"anchor_text":"乙方离职后两年内不得从事与甲方相同或类似业务",'
        '"tracked_changes":"乙方离职后[-两年-][+一年+]内不得从事与甲方相同或类似业务",'
        '"clean_version":"乙方离职后一年内不得从事与甲方相同或类似业务",'
        '"reason":"期限过长","legal_basis":"《劳动合同法》相关规定 [待查]",'
        '"risk_level":"高","negotiation_point":"可缩短","verify_status":"pending"}]', encoding="utf-8")
    out = tmp_path / "out.docx"
    res = redline_docx.apply_redline(str(inp), str(plan), str(out))
    assert out.exists() and res["degraded"] is False and res["changes"] >= 1
    xml = zipfile.ZipFile(str(out)).read("word/document.xml").decode("utf-8")
    assert "w:ins" in xml and "w:del" in xml
    assert "一年" in xml and "两年" in xml

def test_delete_and_insert(tmp_path):
    p = tmp_path / "in2.docx"
    d = Document(); d.add_paragraph("第一条 标的。"); d.add_paragraph("第二条 价款。"); d.save(str(p))
    plan = tmp_path / "plan2.json"
    plan.write_text('[{"clause":"第2条","action":"delete","anchor_text":"第二条 价款。",'
        '"clean_version":"","reason":"冗余","legal_basis":"[待查]","risk_level":"中","verify_status":"pending"},'
        '{"clause":"新增","action":"insert","anchor_text":"第一条 标的。",'
        '"clean_version":"第一条之二 质量标准与验收。","reason":"补缺","legal_basis":"[待查]",'
        '"risk_level":"中","verify_status":"pending"}]', encoding="utf-8")
    out = tmp_path / "out2.docx"
    res = redline_docx.apply_redline(str(p), str(plan), str(out))
    xml = zipfile.ZipFile(str(out)).read("word/document.xml").decode("utf-8")
    assert "w:del" in xml and "w:ins" in xml and res["changes"] >= 2

def test_change_carries_comment(tmp_path):
    inp = _make_input(tmp_path)
    plan = tmp_path / "p.json"
    plan.write_text('[{"clause":"第8条","action":"replace",'
        '"anchor_text":"乙方离职后两年内不得从事与甲方相同或类似业务",'
        '"tracked_changes":"乙方离职后[-两年-][+一年+]内不得从事与甲方相同或类似业务",'
        '"clean_version":"乙方离职后一年内不得从事与甲方相同或类似业务",'
        '"reason":"期限过长","legal_basis":"《劳动合同法》相关规定 [待查]",'
        '"risk_level":"高","verify_status":"pending"}]', encoding="utf-8")
    out = tmp_path / "o.docx"
    res = redline_docx.apply_redline(str(inp), str(plan), str(out))
    assert res["comments"] >= 1
    names = zipfile.ZipFile(str(out)).namelist()
    assert "word/comments.xml" in names
    assert "期限过长" in zipfile.ZipFile(str(out)).read("word/comments.xml").decode("utf-8")

def test_comment_action_no_del(tmp_path):
    inp = _make_input(tmp_path)
    plan = tmp_path / "p2.json"
    plan.write_text('[{"clause":"第8条","action":"comment",'
        '"anchor_text":"乙方离职后两年内不得从事与甲方相同或类似业务","clean_version":"",'
        '"reason":"提示商务确认","legal_basis":"[待查]","risk_level":"低","verify_status":"pending"}]', encoding="utf-8")
    out = tmp_path / "o2.docx"
    res = redline_docx.apply_redline(str(inp), str(plan), str(out))
    xml = zipfile.ZipFile(str(out)).read("word/document.xml").decode("utf-8")
    assert res["comments"] >= 1 and "w:del" not in xml

def test_replace_preserves_surrounding_text(tmp_path):
    p = tmp_path / "in.docx"
    d = Document(); d.add_paragraph("第8条 竞业限制:乙方离职后两年内不得从事与甲方相同或类似业务。"); d.save(str(p))
    plan = tmp_path / "plan.json"
    plan.write_text('[{"clause":"第8条","action":"replace",'
        '"anchor_text":"乙方离职后两年内不得从事与甲方相同或类似业务",'
        '"tracked_changes":"乙方离职后[-两年-][+一年+]内不得从事与甲方相同或类似业务",'
        '"clean_version":"乙方离职后一年内不得从事与甲方相同或类似业务",'
        '"reason":"期限过长","legal_basis":"[待查]","risk_level":"高","verify_status":"pending"}]', encoding="utf-8")
    out = tmp_path / "out.docx"
    import redline_docx; redline_docx.apply_redline(str(p), str(plan), str(out))
    import zipfile
    xml = zipfile.ZipFile(str(out)).read("word/document.xml").decode("utf-8")
    # surrounding text must survive:
    assert "第8条 竞业限制" in xml and "业务。" in xml
    # and the tracked change still applied:
    assert "w:ins" in xml and "w:del" in xml and "一年" in xml

def test_degrade_to_markdown(tmp_path, monkeypatch):
    monkeypatch.setattr(redline_docx, "DOCX_AVAILABLE", False)
    inp = tmp_path / "in.txt"; inp.write_text("乙方离职后两年内不得...", encoding="utf-8")
    plan = tmp_path / "p3.json"
    plan.write_text('[{"clause":"第8条","action":"replace","anchor_text":"乙方离职后两年内不得...",'
        '"clean_version":"乙方离职后一年内不得...","reason":"期限过长","legal_basis":"[待查]",'
        '"risk_level":"高","verify_status":"pending"}]', encoding="utf-8")
    out = tmp_path / "o.docx"
    res = redline_docx.apply_redline(str(inp), str(plan), str(out))
    assert res["degraded"] is True
    md = out.with_suffix(".md")
    assert md.exists() and "期限过长" in md.read_text(encoding="utf-8")
